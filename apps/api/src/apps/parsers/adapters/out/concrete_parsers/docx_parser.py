"""
DOCX парсер для извлечения текста и изображений из .docx файлов.

Использует библиотеку python-docx для извлечения текста, таблиц и структуры документа.
"""
import logging
from io import BytesIO
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from ....domain.ports import DocumentParser
from ....domain.models import ParsedDocument, DocumentImage

logger = logging.getLogger(__name__)


class DocxDocumentParser(DocumentParser):
    """
    Парсер для DOCX файлов.
    
    Извлекает:
    - Текст из параграфов, таблиц и других элементов
    - Структуру документа (заголовки)
    - Изображения (TODO: будет реализовано позже)
    """

    def __init__(
        self,
        min_width: int = 50,
        min_height: int = 50,
        min_file_size: int = 3 * 1024,
    ):
        """
        Args:
            min_width: Минимальная ширина изображения в пикселях
            min_height: Минимальная высота изображения в пикселях
            min_file_size: Минимальный размер файла изображения в байтах
        """
        self.min_width = min_width
        self.min_height = min_height
        self.min_file_size = min_file_size

    def parse(
        self,
        file_bytes: bytes,
        file_name: str,
        process_images: bool = False,
    ) -> ParsedDocument:
        """
        Парсит DOCX документ.

        Args:
            file_bytes: Байты DOCX файла
            file_name: Имя файла (для логирования)
            process_images: Извлекать ли изображения (TODO: пока не реализовано)

        Returns:
            ParsedDocument с результатом парсинга

        Raises:
            Exception: Если не удается обработать DOCX файл
        """
        try:
            logger.info(f"📄 Парсинг DOCX: {file_name}")
            
            # Открываем документ из байтов
            doc_stream = BytesIO(file_bytes)
            doc = Document(doc_stream)
            
            logger.info(f"DOCX открыт. Количество параграфов: {len(doc.paragraphs)}")
            
            
            
            # Извлекаем текст из документа
            text_parts = self.extract_text_from_document(doc)
            final_text = "\n\n".join(text_parts)
            
            # Изображения
            images: list[DocumentImage] = []
            if process_images:
                logger.warning(
                    "Извлечение изображений из DOCX пока не реализовано. "
                    "Это будет добавлено в будущих версиях."
                )
                # TODO: Реализовать извлечение изображений
                # Можно использовать doc.inline_shapes и doc.part.rels для доступа к изображениям
           
            
            return ParsedDocument(
                text=final_text,
                images=images,
                contents=[],
                is_book=False,  # DOCX обычно не являются книгами
            )
            
        except Exception as e:
            logger.error(f"❌ Ошибка при парсинге DOCX файла {file_name}: {str(e)}")
            raise Exception(f"Ошибка при парсинге DOCX файла: {str(e)}")
    
    
    def extract_text_from_document(self, doc: Document) -> list[str]:
        """
        Извлекает текст из документа DOCX.
        
        Обрабатывает параграфы и таблицы в порядке их появления в документе.
        
        Args:
            doc: Объект Document из python-docx
            
        Returns:
            Список текстовых блоков
        """
        text_parts = []
        
        # Обходим все элементы документа (параграфы и таблицы)
        for element in doc.element.body:
            # Проверяем тип элемента
            if element.tag.endswith('p'):  # Параграф
                # Находим соответствующий объект Paragraph
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:
                            text_parts.append(text)
                        break
            
            elif element.tag.endswith('tbl'):  # Таблица
                # Находим соответствующий объект Table
                for table in doc.tables:
                    if table._element == element:
                        table_text = self.format_table(table)
                        if table_text:
                            text_parts.append(table_text)
                        break
        
        return text_parts
    
    def format_table(self, table: Table) -> str:
        """
        Форматирует таблицу в текстовый вид.
        
        Args:
            table: Объект Table из python-docx
            
        Returns:
            Отформатированная таблица
        """
        if not table.rows:
            return ""
        
        # Собираем все строки таблицы
        table_rows = []
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells]
            table_rows.append(cell_texts)
        
        if not table_rows:
            return ""
        
        # Вычисляем ширину колонок
        num_cols = max(len(row) for row in table_rows)
        col_widths = [0] * num_cols
        
        for row in table_rows:
            for i, cell in enumerate(row):
                if i < num_cols:
                    col_widths[i] = max(col_widths[i], len(cell))
        
        # Форматируем таблицу
        formatted_rows = []
        for i, row in enumerate(table_rows):
            # Дополняем строку пустыми ячейками если нужно
            while len(row) < num_cols:
                row.append("")
            
            # Форматируем ячейки с выравниванием
            formatted_cells = [
                cell.ljust(col_widths[j]) for j, cell in enumerate(row)
            ]
            formatted_rows.append("| " + " | ".join(formatted_cells) + " |")
            
            # Добавляем разделитель после заголовка (первой строки)
            if i == 0:
                separator = "|" + "|".join(["-" * (w + 2) for w in col_widths]) + "|"
                formatted_rows.append(separator)
        
        return "\n".join(formatted_rows)
